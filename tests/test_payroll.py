"""Payroll module tests (spec-conformant: UUID, companies/employees, array
components with percent_of, /api/v1/enterprise/payroll).

Infra notes (unchanged rationale):
- psycopg2 (sync) for fixture truncation/seeding so the asyncpg pool is untouched.
- Each integration test uses its own TestClient context so the anyio portal
  (and asyncpg pool) is recreated per test (Windows ProactorEventLoop safety).
- compute_payslip is pure Decimal math, called directly in unit tests.
"""
import contextlib
import os
from decimal import Decimal
from typing import Any, Generator, Iterator

import psycopg2
import pytest
from fastapi import status
from fastapi.testclient import TestClient

from app.constants import PayrollCycleStatus, PayslipStatus
from app.core.security import hash_password
from app.main import app
from app.models.payroll import SalaryStructure
from app.services.payroll_service import compute_payslip

# DSN is env-overridable: set PAYROLL_TEST_DSN. Defaults to a DEDICATED scratch
# database (`payroll_scratch`), SEPARATE from the dev app's DB (`payroll_test`),
# so the destructive fixtures below can never wipe real data. Create/migrate it
# with: `python scripts/setup_test_db.py` (or DB_NAME=payroll_scratch alembic
# upgrade head).
_DSN = os.getenv(
    "PAYROLL_TEST_DSN",
    "host=localhost port=5432 dbname=payroll_scratch user=postgres password=mysql",
)

# Safety guard (belt-and-suspenders with the separate DB above). The fixtures
# TRUNCATE tables, so require an explicit opt-in: set PAYROLL_ALLOW_DB_WIPE=1.
# A bare `pytest` skips them. NEVER point PAYROLL_TEST_DSN at the dev DB.
_ALLOW_DB_WIPE = os.getenv("PAYROLL_ALLOW_DB_WIPE") == "1"


def _require_wipe_allowed() -> None:
    # Hard stop: never let the destructive fixtures run against the dev database,
    # whatever the flags say. This is the guarantee that test runs can't wipe
    # data you entered through the app.
    if "payroll_test" in _DSN:
        raise RuntimeError(
            "Refusing to run destructive tests against the dev database "
            "'payroll_test'. Point PAYROLL_TEST_DSN at the scratch DB "
            "('payroll_scratch') instead."
        )
    if not _ALLOW_DB_WIPE:
        pytest.skip(
            "Destructive DB fixtures are disabled. Set PAYROLL_ALLOW_DB_WIPE=1 "
            "to run them (defaults to the dedicated 'payroll_scratch' DB).",
        )

COMPANY_ID = "00000000-0000-0000-0000-000000000001"
EMP1 = "11111111-1111-1111-1111-111111111111"  # has a salary structure
EMP2 = "22222222-2222-2222-2222-222222222222"  # no structure -> skipped

ADMIN = ("admin@croar.com", "admin123")
VIEWER = ("viewer@croar.com", "viewer123")


def _reset_db() -> None:
    conn = psycopg2.connect(_DSN)
    conn.autocommit = True
    cur = conn.cursor()
    cur.execute("TRUNCATE TABLE tds_challans CASCADE;")
    cur.execute("TRUNCATE TABLE employee_tax_profiles CASCADE;")
    cur.execute("TRUNCATE TABLE payslips CASCADE;")
    cur.execute("TRUNCATE TABLE salary_structures CASCADE;")
    cur.execute("TRUNCATE TABLE salary_templates CASCADE;")
    cur.execute("TRUNCATE TABLE payroll_cycles CASCADE;")
    cur.execute("TRUNCATE TABLE employees CASCADE;")
    # Ensure the default scoping company exists.
    cur.execute(
        "INSERT INTO companies (id, name, currency, created_at, updated_at) "
        "VALUES (%s, 'Croar Technologies', 'INR', now(), now()) "
        "ON CONFLICT (id) DO NOTHING;",
        (COMPANY_ID,),
    )
    for emp_id, first in ((EMP1, "John"), (EMP2, "Jane")):
        cur.execute(
            "INSERT INTO employees (id, company_id, first_name, last_name, email, "
            "created_at, updated_at) VALUES (%s, %s, %s, 'Doe', %s, now(), now());",
            (emp_id, COMPANY_ID, first, f"{first.lower()}@example.com"),
        )
    # Ensure auth users exist (idempotent; survives across runs).
    for email, pw, role in (
        (ADMIN[0], ADMIN[1], "ADMIN"),
        (VIEWER[0], VIEWER[1], "VIEWER"),
    ):
        cur.execute(
            "INSERT INTO users (company_id, email, full_name, hashed_password, "
            "role, is_active, created_at, updated_at) VALUES "
            "(%s, %s, %s, %s, %s, true, now(), now()) "
            "ON CONFLICT (company_id, email) DO NOTHING;",
            (COMPANY_ID, email, email, hash_password(pw), role),
        )
    cur.close()
    conn.close()


def _token(client: TestClient, email: str, password: str) -> str:
    resp = client.post(
        "/api/v1/auth/login", json={"email": email, "password": password}
    )
    assert resp.status_code == status.HTTP_200_OK, resp.text
    return resp.json()["access_token"]


def _delete_user_and_company(email: str) -> None:
    """Remove any user with this email and its owning company (FK cascade drops
    the user too). Used to keep self-service signup tests idempotent, since the
    db_setup fixture does not truncate users/companies."""
    conn = psycopg2.connect(_DSN)
    conn.autocommit = True
    cur = conn.cursor()
    cur.execute("SELECT company_id FROM users WHERE email = %s;", (email,))
    for (company_id,) in cur.fetchall():
        if str(company_id) == COMPANY_ID:
            continue  # never drop the shared default company
        cur.execute("DELETE FROM companies WHERE id = %s;", (str(company_id),))
    cur.close()
    conn.close()


@contextlib.contextmanager
def authed_client(creds: tuple[str, str] = ADMIN) -> Iterator[TestClient]:
    """A TestClient with a Bearer token pre-set for the given user."""
    with TestClient(app) as c:
        token = _token(c, *creds)
        c.headers.update({"Authorization": f"Bearer {token}"})
        yield c


@pytest.fixture(autouse=True)
def db_setup() -> Generator[None, None, None]:
    _require_wipe_allowed()  # skips the test if PAYROLL_ALLOW_DB_WIPE != "1"
    _reset_db()
    yield
    conn = psycopg2.connect(_DSN)
    conn.autocommit = True
    cur = conn.cursor()
    cur.execute("TRUNCATE TABLE tds_challans CASCADE;")
    cur.execute("TRUNCATE TABLE employee_tax_profiles CASCADE;")
    cur.execute("TRUNCATE TABLE payslips CASCADE;")
    cur.execute("TRUNCATE TABLE salary_structures CASCADE;")
    cur.execute("TRUNCATE TABLE salary_templates CASCADE;")
    cur.execute("TRUNCATE TABLE payroll_cycles CASCADE;")
    cur.execute("TRUNCATE TABLE employees CASCADE;")
    cur.close()
    conn.close()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _make_struct(**overrides: Any) -> SalaryStructure:
    defaults: dict[str, Any] = {
        "ctc": Decimal("1200000.00"),
        "currency": "INR",
        "components": [
            {"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 40000},
            {"code": "HRA", "label": "HRA", "type": "percent", "percent": 40, "percent_of": "BASIC"},
            {"code": "SPECIAL", "label": "Special Allowance", "type": "fixed", "amount": 15000},
        ],
        "default_deductions": [
            {"code": "PF", "label": "Provident Fund", "type": "fixed", "amount": 1800},
        ],
    }
    defaults.update(overrides)
    return SalaryStructure(**defaults)


def _create_structure(client: TestClient, employee_id: str = EMP1, **overrides: Any) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "employee_id": employee_id,
        "ctc": 1200000.00,
        "currency": "INR",
        "pay_frequency": "MONTHLY",
        "effective_from": "2026-06-01",
        "components": [
            {"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 40000},
            {"code": "HRA", "label": "HRA", "type": "percent", "percent": 40, "percent_of": "BASIC"},
        ],
        "default_deductions": [
            {"code": "PF", "label": "Provident Fund", "type": "fixed", "amount": 1800},
        ],
        "is_active": True,
    }
    payload.update(overrides)
    resp = client.post("/api/v1/enterprise/payroll/structures", json=payload)
    assert resp.status_code == status.HTTP_201_CREATED, resp.text
    return resp.json()


def _create_cycle(client: TestClient, **overrides: Any) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "name": "June 2026",
        "period_start": "2026-06-01",
        "period_end": "2026-06-30",
        "pay_date": "2026-07-01",
        "notes": "Test cycle",
    }
    payload.update(overrides)
    resp = client.post("/api/v1/enterprise/payroll/cycles", json=payload)
    assert resp.status_code == status.HTTP_201_CREATED, resp.text
    return resp.json()


# ---------------------------------------------------------------------------
# Unit tests — compute_payslip (array shape + percent_of)
# ---------------------------------------------------------------------------
def test_compute_basic_hra_percent_of_special() -> None:
    """BASIC 40000 + HRA 40% of BASIC (16000) + SPECIAL 15000 -> gross 71000."""
    res = compute_payslip(_make_struct(), Decimal("0"), Decimal("30"))
    assert res["gross_earnings"] == Decimal("71000.00")
    assert res["total_deductions"] == Decimal("1800.00")
    assert res["net_pay"] == Decimal("69200.00")
    hra = next(e for e in res["earnings"] if e["code"] == "HRA")
    assert hra["amount"] == 16000.00


def test_compute_percent_of_gross_when_omitted() -> None:
    """A percent deduction with no percent_of is a percent of gross."""
    struct = _make_struct(
        components=[{"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 30000}],
        default_deductions=[{"code": "TAX", "label": "Tax", "type": "percent", "percent": 10}],
    )
    res = compute_payslip(struct, Decimal("0"), Decimal("30"))
    assert res["gross_earnings"] == Decimal("30000.00")
    assert res["total_deductions"] == Decimal("3000.00")
    assert res["net_pay"] == Decimal("27000.00")


def test_compute_lop_proration() -> None:
    """LOP 3/30 -> earnings pro-rated by 0.9; fixed deduction not pro-rated."""
    res = compute_payslip(_make_struct(), Decimal("3"), Decimal("30"))
    assert res["paid_days"] == Decimal("27")
    # 36000 + 14400 + 13500 = 63900
    assert res["gross_earnings"] == Decimal("63900.00")
    assert res["total_deductions"] == Decimal("1800.00")
    assert res["net_pay"] == Decimal("62100.00")


def test_compute_rounding_half_up() -> None:
    struct = _make_struct(
        components=[{"code": "BASIC", "label": "Basic", "type": "fixed", "amount": "10000.005"}],
        default_deductions=[],
    )
    res = compute_payslip(struct, Decimal("0"), Decimal("30"))
    assert res["gross_earnings"] == Decimal("10000.01")


def test_compute_invalid_lop() -> None:
    with pytest.raises(ValueError):
        compute_payslip(_make_struct(), Decimal("31"), Decimal("30"))


# ---------------------------------------------------------------------------
# Unit tests — CTC-driven (dynamic) components: percent-of-CTC + balance line
# ---------------------------------------------------------------------------
def _ctc_driven_components() -> list[dict[str, Any]]:
    """A CTC-driven template: BASIC = 40% of CTC, HRA = 50% of BASIC, and a
    SPECIAL 'balance' line that absorbs the remainder of the CTC."""
    return [
        {"code": "BASIC", "label": "Basic", "type": "percent", "percent": 40, "percent_of": "CTC"},
        {"code": "HRA", "label": "HRA", "type": "percent", "percent": 50, "percent_of": "BASIC"},
        {"code": "SPECIAL", "label": "Special Allowance", "type": "balance"},
    ]


def test_compute_percent_of_ctc() -> None:
    """percent_of 'CTC' resolves against per-period CTC (annual / 12)."""
    struct = _make_struct(
        ctc=Decimal("1200000.00"),  # -> 100000 / month
        components=[{"code": "BASIC", "label": "Basic", "type": "percent", "percent": 40, "percent_of": "CTC"}],
        default_deductions=[],
    )
    res = compute_payslip(struct, Decimal("0"), Decimal("30"))
    assert res["gross_earnings"] == Decimal("40000.00")  # 40% of 100000


def test_compute_balance_line_sums_to_ctc() -> None:
    """BASIC 40000 + HRA 20000 + SPECIAL (balance) -> gross == period CTC."""
    struct = _make_struct(
        ctc=Decimal("1200000.00"),  # 100000 / month
        components=_ctc_driven_components(),
        default_deductions=[],
    )
    res = compute_payslip(struct, Decimal("0"), Decimal("30"))
    assert res["gross_earnings"] == Decimal("100000.00")
    special = next(e for e in res["earnings"] if e["code"] == "SPECIAL")
    assert special["amount"] == 40000.00  # 100000 - 40000 - 20000


def test_compute_balance_scales_with_ctc() -> None:
    """The same template applied at a higher CTC scales every line — the
    'dynamic' property: components are rules, not frozen amounts."""
    res = compute_payslip(
        _make_struct(ctc=Decimal("2400000.00"), components=_ctc_driven_components(), default_deductions=[]),
        Decimal("0"),
        Decimal("30"),
    )
    assert res["gross_earnings"] == Decimal("200000.00")
    amounts = {e["code"]: e["amount"] for e in res["earnings"]}
    assert amounts == {"BASIC": 80000.00, "HRA": 40000.00, "SPECIAL": 80000.00}


def test_compute_balance_never_negative() -> None:
    """If fixed lines already exceed CTC, the balance line floors at 0."""
    struct = _make_struct(
        ctc=Decimal("600000.00"),  # 50000 / month
        components=[
            {"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 60000},
            {"code": "SPECIAL", "label": "Special", "type": "balance"},
        ],
        default_deductions=[],
    )
    res = compute_payslip(struct, Decimal("0"), Decimal("30"))
    special = next(e for e in res["earnings"] if e["code"] == "SPECIAL")
    assert special["amount"] == 0.00
    assert res["gross_earnings"] == Decimal("60000.00")


# ---------------------------------------------------------------------------
# Integration tests — salary templates (reusable, CTC-driven)
# ---------------------------------------------------------------------------
_TEMPLATE_BODY = {
    "name": "Engineer L1",
    "description": "Standard engineering package",
    "components": [
        {"code": "BASIC", "label": "Basic", "type": "percent", "percent": 40, "percent_of": "CTC"},
        {"code": "HRA", "label": "HRA", "type": "percent", "percent": 50, "percent_of": "BASIC"},
        {"code": "SPECIAL", "label": "Special Allowance", "type": "balance"},
    ],
    "default_deductions": [],
}


def test_template_crud_and_unique_name() -> None:
    with authed_client() as c:
        created = c.post("/api/v1/enterprise/payroll/templates", json=_TEMPLATE_BODY)
        assert created.status_code == status.HTTP_201_CREATED, created.text
        tid = created.json()["id"]

        # Duplicate name -> 409
        dup = c.post("/api/v1/enterprise/payroll/templates", json=_TEMPLATE_BODY)
        assert dup.status_code == status.HTTP_409_CONFLICT

        listed = c.get("/api/v1/enterprise/payroll/templates")
        assert listed.status_code == status.HTTP_200_OK
        assert len(listed.json()) == 1

        # Update + soft delete frees the name.
        upd = c.put(
            f"/api/v1/enterprise/payroll/templates/{tid}",
            json={"description": "Updated"},
        )
        assert upd.status_code == status.HTTP_200_OK
        assert upd.json()["description"] == "Updated"

        assert c.delete(f"/api/v1/enterprise/payroll/templates/{tid}").status_code == status.HTTP_200_OK
        assert len(c.get("/api/v1/enterprise/payroll/templates").json()) == 0
        # Name is reusable after delete.
        assert c.post("/api/v1/enterprise/payroll/templates", json=_TEMPLATE_BODY).status_code == status.HTTP_201_CREATED


def test_template_apply_scales_per_employee_ctc() -> None:
    """Applying one template to two employees at different CTCs yields structures
    whose amounts scale to each CTC, then a run computes the scaled pay."""
    with authed_client() as c:
        tid = c.post("/api/v1/enterprise/payroll/templates", json=_TEMPLATE_BODY).json()["id"]

        applied = c.post(
            f"/api/v1/enterprise/payroll/templates/{tid}/apply",
            json={
                "assignments": [
                    {"employee_id": EMP1, "ctc": 1200000, "effective_from": "2026-06-01"},
                    {"employee_id": EMP2, "ctc": 2400000, "effective_from": "2026-06-01"},
                ]
            },
        )
        assert applied.status_code == status.HTTP_200_OK, applied.text
        assert len(applied.json()["created"]) == 2
        assert applied.json()["skipped"] == []

        # Both employees now have a structure linked to the template at their CTC.
        for emp, expected_ctc in ((EMP1, 1200000.0), (EMP2, 2400000.0)):
            rows = c.get(f"/api/v1/enterprise/payroll/structures?employee_id={emp}").json()
            assert len(rows) == 1
            assert rows[0]["template_id"] == tid
            assert float(rows[0]["ctc"]) == expected_ctc

        # Run a cycle: EMP1 gross = 100000, EMP2 gross = 200000.
        cid = _create_cycle(c)["id"]
        run = c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")
        assert run.status_code == status.HTTP_200_OK, run.text
        assert run.json()["created"] == 2
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/mark-paid")
        slips = {s["employee_id"]: s for s in c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/payslips").json()}
        assert float(slips[EMP1]["gross_earnings"]) == 100000.0
        assert float(slips[EMP2]["gross_earnings"]) == 200000.0


def test_template_apply_replace_and_skip() -> None:
    with authed_client() as c:
        tid = c.post("/api/v1/enterprise/payroll/templates", json=_TEMPLATE_BODY).json()["id"]
        base = {"assignments": [{"employee_id": EMP1, "ctc": 1200000, "effective_from": "2026-06-01"}]}

        first = c.post(f"/api/v1/enterprise/payroll/templates/{tid}/apply", json=base)
        assert len(first.json()["created"]) == 1

        # replace_existing=False -> skip the already-configured employee.
        skip = c.post(
            f"/api/v1/enterprise/payroll/templates/{tid}/apply",
            json={**base, "replace_existing": False},
        )
        assert skip.json()["created"] == []
        assert len(skip.json()["skipped"]) == 1

        # Default replace -> new active structure, exactly one remains active.
        repl = c.post(f"/api/v1/enterprise/payroll/templates/{tid}/apply", json=base)
        assert len(repl.json()["created"]) == 1
        active = c.get(f"/api/v1/enterprise/payroll/structures?employee_id={EMP1}").json()
        assert sum(1 for s in active if s["is_active"]) == 1


# ---------------------------------------------------------------------------
# Integration tests
# ---------------------------------------------------------------------------
def test_payroll_lifecycle() -> None:
    """DRAFT -> run -> PROCESSING -> approve -> mark-paid -> PAID."""
    with authed_client() as c:
        _create_structure(c, employee_id=EMP1)

        # Duplicate active structure -> 409
        dup = c.post(
            "/api/v1/enterprise/payroll/structures",
            json={
                "employee_id": EMP1, "ctc": 1200000.0, "effective_from": "2026-06-01",
                "components": [{"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 40000}],
                "default_deductions": [], "is_active": True,
            },
        )
        assert dup.status_code == status.HTTP_409_CONFLICT

        cycle = _create_cycle(c)
        cid = cycle["id"]
        assert cycle["status"] == PayrollCycleStatus.DRAFT

        run = c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")
        assert run.status_code == status.HTTP_200_OK, run.text
        body = run.json()
        assert body["created"] == 1
        assert body["updated"] == 0
        assert len(body["skipped"]) == 1
        assert body["skipped"][0]["employee_id"] == EMP2

        got = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}").json()
        assert got["status"] == PayrollCycleStatus.PROCESSING
        # BASIC 40000 + HRA 40% of BASIC 16000 = 56000 gross; PF 1800 fixed
        assert got["totals"]["headcount"] == 1
        assert got["totals"]["gross"] == 56000.0
        assert got["totals"]["deductions"] == 1800.0
        assert got["totals"]["net"] == 54200.0

        # The computed payslip summary is visible once a run generates it
        # (PROCESSING) — names + amounts. The full individual payslip stays
        # gated to PAID (checked below).
        processing_slips = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/payslips").json()
        assert len(processing_slips) == 1
        assert processing_slips[0]["employee_id"] == EMP1
        assert float(processing_slips[0]["gross_earnings"]) == 56000.0

        # Re-run on PROCESSING is allowed (idempotent refresh)
        rerun = c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")
        assert rerun.status_code == status.HTTP_200_OK
        assert rerun.json()["updated"] == 1

        approve = c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")
        assert approve.status_code == status.HTTP_200_OK
        assert approve.json()["status"] == PayrollCycleStatus.APPROVED

        # Summary still listed while APPROVED, but a direct payslip fetch is
        # forbidden pre-PAID (the full document is released only on disbursement).
        assert len(c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/payslips").json()) == 1
        conn = psycopg2.connect(_DSN)
        conn.autocommit = True
        cur = conn.cursor()
        cur.execute("SELECT id FROM payslips WHERE cycle_id = %s;", (cid,))
        hidden_slip_id = cur.fetchone()[0]
        cur.close()
        conn.close()
        assert (
            c.get(f"/api/v1/enterprise/payroll/payslips/{hidden_slip_id}").status_code
            == status.HTTP_403_FORBIDDEN
        )

        # Run after approve -> 409 (locked)
        assert c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run").status_code == status.HTTP_409_CONFLICT

        paid = c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/mark-paid")
        assert paid.status_code == status.HTTP_200_OK
        assert paid.json()["status"] == PayrollCycleStatus.PAID

        slips = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/payslips").json()
        assert len(slips) == 1
        assert slips[0]["status"] == PayslipStatus.PAID
        assert slips[0]["paid_at"] is not None

        detail = c.get(f"/api/v1/enterprise/payroll/payslips/{slips[0]['id']}").json()
        codes = {line["code"] for line in detail["earnings"]}
        assert codes == {"BASIC", "HRA"}


def test_run_uses_structure_lop_days() -> None:
    """HR sets lop_days on the salary structure; run pro-rates earnings by it.

    Pins the company to the fixed 30-day basis (calendar-derived working days
    OFF) so the proration is tested in isolation from the work calendar."""
    with authed_client() as c:
        c.put(
            "/api/v1/enterprise/calendar/config",
            json={"use_calendar_working_days": False},
        )
        # BASIC 40000 + HRA 40% = 56000 gross; 3 LOP days over 30 -> x0.9
        _create_structure(c, employee_id=EMP1, lop_days=3)
        cid = _create_cycle(c, name="LOP cycle")["id"]
        assert c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run").status_code == status.HTTP_200_OK

        got = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}").json()
        # 36000 + 14400 = 50400 gross; PF 1800 fixed -> 48600 net
        assert got["totals"]["gross"] == 50400.0
        assert got["totals"]["net"] == 48600.0

        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/mark-paid")
        slip = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/payslips").json()[0]
        assert Decimal(str(slip["lop_days"])) == Decimal("3")
        assert Decimal(str(slip["paid_days"])) == Decimal("27")


def test_run_idempotent_on_draft() -> None:
    with authed_client() as c:
        _create_structure(c, employee_id=EMP1)
        cid = _create_cycle(c, name="Idempotency")["id"]
        r1 = c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run").json()
        assert r1["created"] == 1 and r1["updated"] == 0

    # Reset to DRAFT and re-run -> updates, no duplicate
    conn = psycopg2.connect(_DSN)
    conn.autocommit = True
    cur = conn.cursor()
    cur.execute("UPDATE payroll_cycles SET status = 'DRAFT' WHERE id = %s;", (cid,))
    cur.close()
    conn.close()

    with authed_client() as c:
        r2 = c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run").json()
        assert r2["created"] == 0 and r2["updated"] == 1


def test_soft_delete_structure() -> None:
    with authed_client() as c:
        created = _create_structure(c, employee_id=EMP1)
        sid = created["id"]
        assert c.delete(f"/api/v1/enterprise/payroll/structures/{sid}").status_code == status.HTTP_200_OK
        assert len(c.get("/api/v1/enterprise/payroll/structures").json()) == 0
        assert c.get(f"/api/v1/enterprise/payroll/structures/{sid}").status_code == status.HTTP_404_NOT_FOUND
        # Re-create after delete works (partial unique index excludes deleted)
        assert _create_structure(c, employee_id=EMP1)["id"] != sid


def test_structures_filter_by_employee() -> None:
    with authed_client() as c:
        _create_structure(c, employee_id=EMP1)
        all_rows = c.get("/api/v1/enterprise/payroll/structures").json()
        assert len(all_rows) == 1
        filtered = c.get(f"/api/v1/enterprise/payroll/structures?employee_id={EMP2}").json()
        assert filtered == []


def test_cancel_and_soft_delete_cycle() -> None:
    with authed_client() as c:
        cid = _create_cycle(c)["id"]
        cancelled = c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/cancel")
        assert cancelled.status_code == status.HTTP_200_OK
        assert cancelled.json()["status"] == PayrollCycleStatus.CANCELLED

        # Soft-delete (not PAID) -> 200, then hidden from list
        cid2 = _create_cycle(c, name="Delete me")["id"]
        assert c.delete(f"/api/v1/enterprise/payroll/cycles/{cid2}").status_code == status.HTTP_200_OK
        ids = {row["id"] for row in c.get("/api/v1/enterprise/payroll/cycles").json()}
        assert cid2 not in ids


def test_invalid_transitions() -> None:
    with authed_client() as c:
        cid = _create_cycle(c)["id"]
        # approve/mark-paid from DRAFT -> 409
        assert c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve").status_code == status.HTTP_409_CONFLICT
        assert c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/mark-paid").status_code == status.HTTP_409_CONFLICT


# ---------------------------------------------------------------------------
# Employees
# ---------------------------------------------------------------------------
def test_employee_crud_and_soft_delete() -> None:
    with authed_client() as c:
        rows = c.get("/api/v1/enterprise/employees").json()
        assert len(rows) == 2  # seeded EMP1, EMP2

        new = c.post(
            "/api/v1/enterprise/employees",
            json={"first_name": "Sam", "last_name": "Lee", "email": "sam@example.com"},
        )
        assert new.status_code == status.HTTP_201_CREATED
        new_id = new.json()["id"]
        assert len(c.get("/api/v1/enterprise/employees").json()) == 3

        # Duplicate email in same company -> 409
        dup = c.post(
            "/api/v1/enterprise/employees",
            json={"first_name": "Sam2", "email": "sam@example.com"},
        )
        assert dup.status_code == status.HTTP_409_CONFLICT

        assert c.delete(f"/api/v1/enterprise/employees/{new_id}").status_code == status.HTTP_204_NO_CONTENT
        ids = {e["id"] for e in c.get("/api/v1/enterprise/employees").json()}
        assert new_id not in ids

        # The email is still held by the soft-deleted row (unique constraint has
        # no deleted_at filter), so re-creating it returns a clean 409, not 500.
        readd = c.post(
            "/api/v1/enterprise/employees",
            json={"first_name": "Sam3", "email": "sam@example.com"},
        )
        assert readd.status_code == status.HTTP_409_CONFLICT

        # Updating one employee's email to another's is also blocked.
        clash = c.put(
            f"/api/v1/enterprise/employees/{EMP1}",
            json={"email": "jane@example.com"},  # seeded EMP2's email
        )
        assert clash.status_code == status.HTTP_409_CONFLICT


# ---------------------------------------------------------------------------
# Auth & RBAC (spec §7)
# ---------------------------------------------------------------------------
def test_login_success_and_permissions() -> None:
    with TestClient(app) as c:
        resp = c.post(
            "/api/v1/auth/login",
            json={"email": ADMIN[0], "password": ADMIN[1]},
        )
        assert resp.status_code == status.HTTP_200_OK, resp.text
        body = resp.json()
        assert body["token_type"] == "bearer"
        assert body["user"]["role"] == "ADMIN"
        assert "payroll:run" in body["user"]["permissions"]


def test_login_bad_credentials() -> None:
    with TestClient(app) as c:
        resp = c.post(
            "/api/v1/auth/login",
            json={"email": ADMIN[0], "password": "wrong"},
        )
        assert resp.status_code == status.HTTP_401_UNAUTHORIZED


def test_protected_routes_require_auth() -> None:
    """No token -> 401 on protected endpoints."""
    with TestClient(app) as c:
        assert c.get("/api/v1/enterprise/payroll/cycles").status_code == status.HTTP_401_UNAUTHORIZED
        assert c.get("/api/v1/enterprise/employees").status_code == status.HTTP_401_UNAUTHORIZED


def test_viewer_can_read_but_not_write() -> None:
    """A VIEWER may list but cannot run payroll or create resources (403)."""
    with authed_client(VIEWER) as c:
        # read is allowed
        assert c.get("/api/v1/enterprise/payroll/cycles").status_code == status.HTTP_200_OK
        # create cycle requires payroll:configure -> 403
        bad = c.post(
            "/api/v1/enterprise/payroll/cycles",
            json={
                "name": "Nope",
                "period_start": "2026-06-01",
                "period_end": "2026-06-30",
                "pay_date": "2026-07-01",
            },
        )
        assert bad.status_code == status.HTTP_403_FORBIDDEN


def test_viewer_cannot_run_cycle() -> None:
    """Viewer is blocked from the run action even on an existing cycle."""
    with authed_client(ADMIN) as admin:
        cid = _create_cycle(admin, name="Locked")["id"]
    with authed_client(VIEWER) as viewer:
        resp = viewer.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")
        assert resp.status_code == status.HTTP_403_FORBIDDEN


def test_users_admin_only() -> None:
    """User administration requires users:manage (ADMIN); viewer gets 403."""
    with authed_client(ADMIN) as c:
        assert c.get("/api/v1/auth/users").status_code == status.HTTP_200_OK
    with authed_client(VIEWER) as c:
        assert c.get("/api/v1/auth/users").status_code == status.HTTP_403_FORBIDDEN


def test_signup_creates_org_and_admin() -> None:
    """Public signup provisions a new company + its first user as ADMIN, returns
    a working token, and rejects a duplicate email."""
    email = "founder@newco-signup.com"
    _delete_user_and_company(email)  # in case a prior run was interrupted
    try:
        with TestClient(app) as c:
            resp = c.post(
                "/api/v1/auth/signup",
                json={
                    "company_name": "NewCo Industries",
                    "full_name": "Fran Founder",
                    "email": email,
                    "password": "secret123",
                },
            )
            assert resp.status_code == status.HTTP_201_CREATED, resp.text
            body = resp.json()
            assert body["user"]["role"] == "ADMIN"
            assert body["user"]["company_id"] != COMPANY_ID  # a brand-new tenant
            assert "users:manage" in body["user"]["permissions"]

            # The returned token authenticates against a protected endpoint.
            me = c.get(
                "/api/v1/auth/me",
                headers={"Authorization": f"Bearer {body['access_token']}"},
            )
            assert me.status_code == status.HTTP_200_OK
            assert me.json()["email"] == email

            # Email is globally unique (login resolves by email alone) -> 409.
            dup = c.post(
                "/api/v1/auth/signup",
                json={
                    "company_name": "Another Co",
                    "full_name": "Imposter",
                    "email": email,
                    "password": "secret123",
                },
            )
            assert dup.status_code == status.HTTP_409_CONFLICT
    finally:
        _delete_user_and_company(email)


# ---------------------------------------------------------------------------
# Per-run adjustments (bonuses / arrears / one-time deductions)
# ---------------------------------------------------------------------------
def _add_adjustment(c: TestClient, cid: str, **body: Any) -> Any:
    return c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/adjustments", json=body)


def test_run_applies_earning_and_deduction_adjustments() -> None:
    """A bonus (earning) adds to gross/net; a recovery (deduction) subtracts.
    Both surface as snapshot lines on the payslip."""
    with authed_client() as c:
        _create_structure(c, employee_id=EMP1)  # gross 56000, deductions 1800
        cid = _create_cycle(c)["id"]

        assert _add_adjustment(
            c, cid, employee_id=EMP1, kind="earning",
            code="BONUS", label="Festival Bonus", amount=5000,
        ).status_code == status.HTTP_201_CREATED
        assert _add_adjustment(
            c, cid, employee_id=EMP1, kind="deduction",
            code="RECOVERY", label="Advance Recovery", amount=1000,
        ).status_code == status.HTTP_201_CREATED

        assert c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run").status_code == status.HTTP_200_OK
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/mark-paid")

        slips = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/payslips").json()
        slip = next(s for s in slips if s["employee_id"] == EMP1)
        assert float(slip["gross_earnings"]) == 61000.0   # 56000 + 5000 bonus
        assert float(slip["total_deductions"]) == 2800.0  # 1800 + 1000 recovery
        assert float(slip["net_pay"]) == 58200.0

        detail = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}").json()
        assert "BONUS" in {e["code"] for e in detail["earnings"]}
        assert "RECOVERY" in {d["code"] for d in detail["deductions"]}


def test_adjustments_listed_and_locked_after_approve() -> None:
    with authed_client() as c:
        _create_structure(c, employee_id=EMP1)
        cid = _create_cycle(c)["id"]
        add = _add_adjustment(
            c, cid, employee_id=EMP1, kind="earning", code="BONUS", label="Bonus", amount=2000
        )
        assert add.status_code == status.HTTP_201_CREATED, add.text
        adj_id = add.json()["id"]

        listed = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/adjustments").json()
        assert [a["id"] for a in listed] == [adj_id]

        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")
        # Approved cycle is locked: no further add or delete.
        assert _add_adjustment(
            c, cid, employee_id=EMP1, kind="earning", code="X", label="X", amount=1
        ).status_code == status.HTTP_409_CONFLICT
        assert c.delete(
            f"/api/v1/enterprise/payroll/adjustments/{adj_id}"
        ).status_code == status.HTTP_409_CONFLICT


def test_delete_adjustment_excludes_it_from_run() -> None:
    with authed_client() as c:
        _create_structure(c, employee_id=EMP1)
        cid = _create_cycle(c)["id"]
        adj_id = _add_adjustment(
            c, cid, employee_id=EMP1, kind="earning", code="BONUS", label="Bonus", amount=5000
        ).json()["id"]
        assert c.delete(
            f"/api/v1/enterprise/payroll/adjustments/{adj_id}"
        ).status_code == status.HTTP_200_OK

        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/mark-paid")
        slips = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/payslips").json()
        slip = next(s for s in slips if s["employee_id"] == EMP1)
        assert float(slip["gross_earnings"]) == 56000.0  # bonus removed before run


# ---------------------------------------------------------------------------
# Reports & registers
# ---------------------------------------------------------------------------
def test_salary_register_and_payroll_summary_reports() -> None:
    with authed_client() as c:
        _create_structure(c, employee_id=EMP1)
        cid = _create_cycle(c)["id"]

        # No payslips yet (DRAFT) -> register is 409.
        assert c.get(
            f"/api/v1/enterprise/reports/salary-register?cycle_id={cid}"
        ).status_code == status.HTTP_409_CONFLICT

        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")

        # Salary register CSV (available at PROCESSING — internal HR report).
        csv_resp = c.get(
            f"/api/v1/enterprise/reports/salary-register?cycle_id={cid}&format=csv"
        )
        assert csv_resp.status_code == status.HTTP_200_OK
        assert csv_resp.headers["content-type"].startswith("text/csv")
        body = csv_resp.content.decode("utf-8-sig")
        assert "Employee Name" in body and "Net Pay" in body
        assert "John" in body  # EMP1's first name from the seed

        # Salary register PDF.
        pdf_resp = c.get(
            f"/api/v1/enterprise/reports/salary-register?cycle_id={cid}&format=pdf"
        )
        assert pdf_resp.status_code == status.HTTP_200_OK
        assert pdf_resp.headers["content-type"] == "application/pdf"
        assert pdf_resp.content[:4] == b"%PDF"

        # Payroll summary CSV lists the cycle.
        sum_resp = c.get("/api/v1/enterprise/reports/payroll-summary?format=csv")
        assert sum_resp.status_code == status.HTTP_200_OK
        assert "Cycle" in sum_resp.content.decode("utf-8-sig")


# ---------------------------------------------------------------------------
# Settings — organisation profile
# ---------------------------------------------------------------------------
def test_organization_profile_get_and_update() -> None:
    with authed_client() as c:  # ADMIN
        org = c.get("/api/v1/enterprise/settings/organization")
        assert org.status_code == status.HTTP_200_OK, org.text
        assert org.json()["id"] == COMPANY_ID

        upd = c.put(
            "/api/v1/enterprise/settings/organization",
            json={
                "name": "Croar Technologies Pvt Ltd",
                "industry": "Information Technology",
                "city": "Bengaluru",
                "country": "India",
                "pan": "aaapz1234c",  # lower-case -> normalised to upper
                "contact_email": "",  # blank -> stored as null
            },
        )
        assert upd.status_code == status.HTTP_200_OK, upd.text
        body = upd.json()
        assert body["name"] == "Croar Technologies Pvt Ltd"
        assert body["industry"] == "Information Technology"
        assert body["pan"] == "AAAPZ1234C"
        assert body["contact_email"] is None

        # Persisted.
        again = c.get("/api/v1/enterprise/settings/organization").json()
        assert again["city"] == "Bengaluru"
        assert again["pan"] == "AAAPZ1234C"


def test_organization_update_is_admin_only() -> None:
    with authed_client(VIEWER) as c:
        assert c.get("/api/v1/enterprise/settings/organization").status_code == status.HTTP_200_OK
        resp = c.put(
            "/api/v1/enterprise/settings/organization", json={"name": "Hacked"}
        )
        assert resp.status_code == status.HTTP_403_FORBIDDEN


def test_payslip_settings_get_update_and_defaults() -> None:
    # NOTE: the scratch DB never truncates `companies`, so this test must not
    # assume a pristine row — it sets its own known baseline first.
    with authed_client() as c:  # ADMIN
        # GET always carries the real company name (UI display fallback).
        assert c.get("/api/v1/enterprise/settings/payslip").json()["company_name"]

        # Baseline: blank text (-> null) + all sections on.
        base = c.put(
            "/api/v1/enterprise/settings/payslip",
            json={
                "display_name": "",
                "logo_url": "",
                "accent_color": "",
                "footer_note": "",
                "show_employer_contributions": True,
                "show_tax_block": True,
                "show_attendance": True,
            },
        ).json()
        assert base["display_name"] is None
        assert base["accent_color"] is None
        assert base["show_attendance"] is True

        # Partial update: branding + one toggle off; others must be untouched.
        b = c.put(
            "/api/v1/enterprise/settings/payslip",
            json={
                "display_name": "ACME Payroll",
                "accent_color": "#16A34A",  # normalised to lower-case
                "footer_note": "  ",  # blank -> null
                "show_attendance": False,
            },
        ).json()
        assert b["display_name"] == "ACME Payroll"
        assert b["accent_color"] == "#16a34a"
        assert b["footer_note"] is None
        assert b["show_attendance"] is False
        assert b["show_tax_block"] is True  # untouched toggle preserved

        # Persisted across requests.
        again = c.get("/api/v1/enterprise/settings/payslip").json()
        assert again["display_name"] == "ACME Payroll"
        assert again["show_attendance"] is False

        # Invalid colour is rejected.
        bad = c.put("/api/v1/enterprise/settings/payslip", json={"accent_color": "green"})
        assert bad.status_code == status.HTTP_422_UNPROCESSABLE_ENTITY


def test_payslip_settings_update_is_admin_only() -> None:
    with authed_client(VIEWER) as c:
        assert c.get("/api/v1/enterprise/settings/payslip").status_code == status.HTTP_200_OK
        resp = c.put(
            "/api/v1/enterprise/settings/payslip", json={"display_name": "Nope"}
        )
        assert resp.status_code == status.HTTP_403_FORBIDDEN


def test_statutory_config_applies_across_payroll() -> None:
    with authed_client() as c:  # ADMIN
        original = c.get("/api/v1/enterprise/settings/statutory").json()
        assert original["pf_wage_ceiling"] == 15000
        assert original["pf_employee_rate"] == 0.12
        try:
            # Override: raise the PF ceiling and lower the employee rate.
            upd = c.put(
                "/api/v1/enterprise/settings/statutory",
                json={"pf_wage_ceiling": 50000, "pf_employee_rate": 0.10},
            )
            assert upd.status_code == status.HTTP_200_OK, upd.text
            b = upd.json()
            assert b["pf_wage_ceiling"] == 50000
            assert b["pf_employee_rate"] == 0.10
            assert b["esi_wage_limit"] == 21000  # untouched default preserved

            # The override flows into the live preview (same engine as a run):
            # PF employee = 10% of min(40000, 50000) = 4000 (was 1800 by default).
            prev = c.post(
                "/api/v1/enterprise/payroll/structures/preview",
                json={
                    "ctc": 1200000,
                    "pay_frequency": "MONTHLY",
                    "components": [{"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 40000}],
                    "default_deductions": [],
                    "lop_days": 0,
                    "pf_enabled": True,
                    "pf_cap_at_ceiling": True,
                    "esi_enabled": False,
                    "pt_enabled": False,
                    "tds_enabled": False,
                },
            ).json()
            pf = [d for d in prev["deductions"] if d["code"] == "PF"]
            assert pf and float(pf[0]["amount"]) == 4000.0

            # Validation: a rate above 100% is rejected.
            bad = c.put("/api/v1/enterprise/settings/statutory", json={"pf_employee_rate": 1.5})
            assert bad.status_code == status.HTTP_422_UNPROCESSABLE_ENTITY
        finally:
            # Restore defaults so other tests compute with standard rates.
            c.put("/api/v1/enterprise/settings/statutory", json=original)


def test_statutory_config_update_is_admin_only() -> None:
    with authed_client(VIEWER) as c:
        assert c.get("/api/v1/enterprise/settings/statutory").status_code == status.HTTP_200_OK
        resp = c.put(
            "/api/v1/enterprise/settings/statutory", json={"pf_employee_rate": 0.10}
        )
        assert resp.status_code == status.HTTP_403_FORBIDDEN


# ---------------------------------------------------------------------------
# Taxes & Forms — IT declarations + TDS challans
# ---------------------------------------------------------------------------
def test_tax_profile_upsert_and_list() -> None:
    with authed_client() as c:
        assert c.get("/api/v1/enterprise/taxes/profiles").json() == []

        r = c.put(
            f"/api/v1/enterprise/taxes/profiles/{EMP1}",
            json={
                "tax_regime": "OLD",
                "declared_80c": 150000,
                "declared_80d": 25000,
                "prev_employer_income": 300000,
            },
        )
        assert r.status_code == status.HTTP_200_OK, r.text
        body = r.json()
        assert body["tax_regime"] == "OLD"
        assert float(body["declared_80c"]) == 150000.0
        assert body["employee_id"] == EMP1

        # Upsert replaces (one profile per employee).
        r2 = c.put(
            f"/api/v1/enterprise/taxes/profiles/{EMP1}",
            json={"tax_regime": "NEW", "declared_80c": 0},
        )
        assert r2.json()["tax_regime"] == "NEW"
        profiles = c.get("/api/v1/enterprise/taxes/profiles").json()
        assert len(profiles) == 1 and profiles[0]["employee_id"] == EMP1


def test_tax_profile_unknown_employee_404() -> None:
    with authed_client() as c:
        r = c.put(
            "/api/v1/enterprise/taxes/profiles/99999999-9999-9999-9999-999999999999",
            json={"tax_regime": "NEW"},
        )
        assert r.status_code == status.HTTP_404_NOT_FOUND


def test_tds_challan_crud() -> None:
    with authed_client() as c:
        assert c.get("/api/v1/enterprise/taxes/challans").json() == []
        r = c.post(
            "/api/v1/enterprise/taxes/challans",
            json={
                "period_month": "2026-04",
                "amount": 50000,
                "challan_number": "CIN12345",
                "bsr_code": "0510308",
                "deposit_date": "2026-05-07",
            },
        )
        assert r.status_code == status.HTTP_201_CREATED, r.text
        challan_id = r.json()["id"]

        listed = c.get("/api/v1/enterprise/taxes/challans").json()
        assert len(listed) == 1 and listed[0]["challan_number"] == "CIN12345"

        assert c.delete(
            f"/api/v1/enterprise/taxes/challans/{challan_id}"
        ).status_code == status.HTTP_200_OK
        assert c.get("/api/v1/enterprise/taxes/challans").json() == []


def test_taxes_are_write_protected_for_viewer() -> None:
    with authed_client(VIEWER) as c:
        assert c.get("/api/v1/enterprise/taxes/challans").status_code == status.HTTP_200_OK
        assert c.post(
            "/api/v1/enterprise/taxes/challans",
            json={
                "period_month": "2026-04",
                "amount": 1,
                "challan_number": "X",
                "deposit_date": "2026-05-07",
            },
        ).status_code == status.HTTP_403_FORBIDDEN
        assert c.put(
            f"/api/v1/enterprise/taxes/profiles/{EMP1}", json={"tax_regime": "NEW"}
        ).status_code == status.HTTP_403_FORBIDDEN


# ---------------------------------------------------------------------------
# Income tax (TDS) engine + integration
# ---------------------------------------------------------------------------
def test_tds_engine_regimes() -> None:
    from app.services import tax_engine

    # New regime: taxable 11.25L < 12L rebate -> nil tax.
    r = tax_engine.compute_tds(annual_gross=Decimal("1200000"), regime="NEW")
    assert float(r["annual_tax"]) == 0.0
    # New regime, higher income -> positive monthly TDS.
    r2 = tax_engine.compute_tds(annual_gross=Decimal("2400000"), regime="NEW")
    assert float(r2["monthly_tds"]) > 0
    # Old regime deductions reduce taxable income.
    base = tax_engine.compute_tds(annual_gross=Decimal("1500000"), regime="OLD")
    with_80c = tax_engine.compute_tds(
        annual_gross=Decimal("1500000"), regime="OLD", declarations={"declared_80c": 150000}
    )
    assert float(with_80c["taxable_income"]) < float(base["taxable_income"])


_TDS_COMPONENTS = [
    {"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 120000},
    {"code": "HRA", "label": "HRA", "type": "fixed", "amount": 50000},
    {"code": "SPL", "label": "Special", "type": "fixed", "amount": 30000},
]


def test_tds_deducted_on_payslip_when_enabled() -> None:
    with authed_client() as c:
        _create_structure(
            c, employee_id=EMP1, tds_enabled=True,
            components=_TDS_COMPONENTS, default_deductions=[],
        )
        cid = _create_cycle(c)["id"]
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/mark-paid")
        slips = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/payslips").json()
        slip = next(s for s in slips if s["employee_id"] == EMP1)
        detail = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}").json()
        tds_lines = [d for d in detail["deductions"] if d["code"] == "TDS"]
        assert len(tds_lines) == 1 and float(tds_lines[0]["amount"]) > 0
        assert detail["statutory"]["tds"]["regime"] == "NEW"


def test_tds_declaration_changes_payslip() -> None:
    """Switching an employee to OLD regime with 80C changes the TDS amount."""
    with authed_client() as c:
        _create_structure(
            c, employee_id=EMP1, tds_enabled=True,
            components=_TDS_COMPONENTS, default_deductions=[],
        )
        c.put(
            f"/api/v1/enterprise/taxes/profiles/{EMP1}",
            json={"tax_regime": "OLD", "declared_80c": 150000, "declared_home_loan_interest": 200000},
        )
        cid = _create_cycle(c)["id"]
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/mark-paid")
        slips = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/payslips").json()
        detail = c.get(
            f"/api/v1/enterprise/payroll/payslips/{slips[0]['id']}"
        ).json()
        assert detail["statutory"]["tds"]["regime"] == "OLD"
        assert float(detail["statutory"]["tds"]["declared_deductions"]) == 350000.0


def test_tds_liabilities_reconciliation() -> None:
    with authed_client() as c:
        _create_structure(
            c, employee_id=EMP1, tds_enabled=True,
            components=_TDS_COMPONENTS, default_deductions=[],
        )
        cid = _create_cycle(c)["id"]  # period June 2026
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")

        libs = c.get("/api/v1/enterprise/taxes/tds-liabilities").json()
        june = next(row for row in libs if row["period_month"] == "2026-06")
        assert float(june["tds_deducted"]) > 0
        assert float(june["tds_deposited"]) == 0.0

        c.post(
            "/api/v1/enterprise/taxes/challans",
            json={
                "period_month": "2026-06",
                "amount": june["tds_deducted"],
                "challan_number": "RECON1",
                "deposit_date": "2026-07-07",
            },
        )
        libs2 = c.get("/api/v1/enterprise/taxes/tds-liabilities").json()
        june2 = next(row for row in libs2 if row["period_month"] == "2026-06")
        assert float(june2["difference"]) == 0.0


# ---------------------------------------------------------------------------
# Audit trail
# ---------------------------------------------------------------------------
def _truncate_audit() -> None:
    conn = psycopg2.connect(_DSN)
    conn.autocommit = True
    cur = conn.cursor()
    cur.execute("TRUNCATE TABLE audit_logs;")
    cur.close()
    conn.close()


def test_audit_trail_records_actions_with_actor() -> None:
    _truncate_audit()
    with authed_client() as c:  # ADMIN
        _create_structure(c, employee_id=EMP1)
        cid = _create_cycle(c)["id"]
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")

        entries = c.get("/api/v1/enterprise/audit").json()
        actions = [e["action"] for e in entries]
        assert "Created payroll cycle" in actions
        assert "Ran payroll cycle" in actions
        assert "Approved payroll cycle" in actions
        # Newest first.
        assert entries[0]["created_at"] >= entries[-1]["created_at"]
        # Actor email is snapshotted; status captured.
        approve = next(e for e in entries if e["action"] == "Approved payroll cycle")
        assert approve["actor_email"] == ADMIN[0]
        assert approve["status_code"] == status.HTTP_200_OK


def test_audit_records_denied_actions() -> None:
    """A VIEWER's blocked mutation (403) is still recorded, with the actor."""
    _truncate_audit()
    with authed_client(VIEWER) as c:
        cid_attempt = c.post(
            "/api/v1/enterprise/payroll/cycles",
            json={
                "name": "Nope",
                "period_start": "2026-06-01",
                "period_end": "2026-06-30",
                "pay_date": "2026-07-01",
            },
        )
        assert cid_attempt.status_code == status.HTTP_403_FORBIDDEN

    with authed_client() as c:  # ADMIN reads the trail
        entries = c.get("/api/v1/enterprise/audit").json()
        denied = [e for e in entries if e["status_code"] == status.HTTP_403_FORBIDDEN]
        assert any(e["actor_email"] == VIEWER[0] for e in denied)


# ---------------------------------------------------------------------------
# Dashboard
# ---------------------------------------------------------------------------
def test_dashboard_summary() -> None:
    with authed_client() as c:
        # Configure EMP1 and run+approve+pay a cycle so paid totals populate.
        _create_structure(c, employee_id=EMP1)
        cid = _create_cycle(c)["id"]
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/mark-paid")

        d = c.get("/api/v1/enterprise/payroll/dashboard")
        assert d.status_code == status.HTTP_200_OK, d.text
        body = d.json()
        assert body["employees"]["total"] == 2
        assert body["employees"]["configured"] == 1
        assert body["employees"]["missing"] == 1
        assert body["active_structures"] == 1
        assert body["cycles"]["total"] == 1
        assert body["cycles"]["by_status"]["PAID"] == 1
        assert float(body["payroll"]["net_paid"]) == 54200.0
        assert body["payroll"]["payslips_paid"] == 1
        assert len(body["recent_cycles"]) == 1


def test_dashboard_requires_auth() -> None:
    with TestClient(app) as c:
        assert c.get("/api/v1/enterprise/payroll/dashboard").status_code == status.HTTP_401_UNAUTHORIZED


# ---------------------------------------------------------------------------
# Payslip PDF + email
# ---------------------------------------------------------------------------
def _paid_payslip(c: TestClient, employee_id: str = EMP1) -> dict[str, Any]:
    """Drive a cycle to PAID and return the resulting (released) payslip."""
    _create_structure(c, employee_id=employee_id)
    cid = _create_cycle(c)["id"]
    assert c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run").status_code == status.HTTP_200_OK
    c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")
    c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/mark-paid")
    return c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/payslips").json()[0]


def test_payslip_pdf_download() -> None:
    with authed_client() as c:
        slip = _paid_payslip(c)
        r = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/pdf")
        assert r.status_code == status.HTTP_200_OK
        assert r.headers["content-type"] == "application/pdf"
        assert r.content[:4] == b"%PDF"
        assert "attachment" in r.headers.get("content-disposition", "")


def _sample_docx_template() -> bytes:
    """A minimal .docx template with docxtpl tokens, built in-memory."""
    import io

    from docx import Document

    d = Document()
    d.add_heading("{{ company_name }} — Payslip", level=1)
    d.add_paragraph("Employee: {{ employee.name }}")
    d.add_paragraph("Net Pay: {{ net }}")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_payslip_doc_template_upload_fill_and_pdf_fallback() -> None:
    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    with authed_client() as c:
        slip = _paid_payslip(c)
        try:
            # Reject non-docx uploads.
            bad = c.put(
                "/api/v1/enterprise/settings/payslip/document",
                files={"file": ("x.txt", b"not a docx", "text/plain")},
            )
            assert bad.status_code == status.HTTP_422_UNPROCESSABLE_ENTITY

            # Upload a real .docx template.
            up = c.put(
                "/api/v1/enterprise/settings/payslip/document",
                files={"file": ("tpl.docx", _sample_docx_template(), docx_mime)},
            )
            assert up.status_code == status.HTTP_200_OK, up.text
            assert up.json()["has_doc_template"] is True
            assert up.json()["doc_filename"] == "tpl.docx"

            # Download the filled .docx — tokens replaced, valid OOXML (starts PK).
            doc = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/docx")
            assert doc.status_code == status.HTTP_200_OK, doc.text
            assert doc.content[:2] == b"PK"
            assert "attachment" in doc.headers.get("content-disposition", "")
            # Assert tokens were actually FILLED (not just a valid file): the
            # employee's name and net pay must appear, and no raw {{ }} remain.
            import io as _io

            from docx import Document as _Doc

            filled = "\n".join(p.text for p in _Doc(_io.BytesIO(doc.content)).paragraphs)
            assert "John Doe" in filled, filled
            assert "{{" not in filled, f"unfilled token remains: {filled}"

            # Enable the template; the PDF endpoint still returns a PDF, falling
            # back to the built-in layout when no docx->pdf converter is present.
            c.put("/api/v1/enterprise/settings/payslip", json={"use_doc_template": True})
            pdf = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/pdf")
            assert pdf.status_code == status.HTTP_200_OK
            assert pdf.content[:4] == b"%PDF"
        finally:
            # Remove the template (also clears use_doc_template) so other tests
            # generate the standard built-in payslip.
            cleared = c.delete("/api/v1/enterprise/settings/payslip/document")
            assert cleared.json()["has_doc_template"] is False
            assert cleared.json()["use_doc_template"] is False


def test_payslip_doc_endpoints_admin_only() -> None:
    with authed_client(VIEWER) as c:
        docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        up = c.put(
            "/api/v1/enterprise/settings/payslip/document",
            files={"file": ("tpl.docx", _sample_docx_template(), docx_mime)},
        )
        assert up.status_code == status.HTTP_403_FORBIDDEN


def test_payslip_docx_404_without_template() -> None:
    with authed_client() as c:
        slip = _paid_payslip(c)
        # No template uploaded for this company in this test's state.
        c.delete("/api/v1/enterprise/settings/payslip/document")
        r = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/docx")
        assert r.status_code == status.HTTP_404_NOT_FOUND


def _blank_payslip_docx() -> bytes:
    """A company's own BLANK payslip: labels + empty value cells, NO tokens —
    exactly what the smart-mapping wizard must handle."""
    import io

    from docx import Document

    d = Document()
    d.add_paragraph("Employee Name: ")
    t = d.add_table(rows=3, cols=4)
    t.cell(0, 0).text = "Basic Salary"; t.cell(0, 2).text = "Provident Fund"
    t.cell(1, 0).text = "House Rent Allowance"; t.cell(1, 2).text = "Professional Tax"
    t.cell(2, 0).text = "Gross Earnings"; t.cell(2, 2).text = "Total Deductions"
    d.add_paragraph("Net Pay: ")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_payslip_smart_mapping_wizard_fills_blank_template() -> None:
    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    with authed_client() as c:
        slip = _paid_payslip(c)
        try:
            # 1) Scan a token-free blank template. Every label→blank pair should
            #    surface, with common labels auto-suggested.
            scan = c.post(
                "/api/v1/enterprise/settings/payslip/document/scan",
                files={"file": ("blank.docx", _blank_payslip_docx(), docx_mime)},
            )
            assert scan.status_code == status.HTTP_200_OK, scan.text
            body = scan.json()
            slots = body["slots"]
            assert len(slots) >= 8, slots
            suggested = {s["label"]: s["suggested_token"] for s in slots}
            assert suggested.get("Basic Salary") == "amount.BASIC"
            assert suggested.get("Net Pay") == "net"
            assert suggested.get("Employee Name") == "employee.name"
            assert any(f["key"] == "amount.HRA" for f in body["fields"])

            # 2) Apply the auto-suggested mapping.
            mapping = {
                str(s["index"]): s["suggested_token"]
                for s in slots
                if s["suggested_token"]
            }
            applied = c.put(
                "/api/v1/enterprise/settings/payslip/document/mapping",
                json={"mapping": mapping},
            )
            assert applied.status_code == status.HTTP_200_OK, applied.text
            assert applied.json()["has_doc_template"] is True
            assert applied.json()["doc_mapped"] is True
            # Applying the mapping turns the doc template on automatically.
            assert applied.json()["use_doc_template"] is True

            # 3) The generated .docx is filled from the mapped template.
            import io as _io

            from docx import Document as _Doc

            doc = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/docx")
            assert doc.status_code == status.HTTP_200_OK, doc.text
            parsed = _Doc(_io.BytesIO(doc.content))
            text = "\n".join(p.text for p in parsed.paragraphs)
            for tbl in parsed.tables:
                for row in tbl.rows:
                    text += "\n" + " | ".join(cell.text for cell in row.cells)
            assert "John Doe" in text, text
            assert "{{" not in text, f"unfilled token remains: {text}"

            # 4) The template drives the on-screen / print HTML view too.
            prev = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/preview-html")
            assert prev.status_code == status.HTTP_200_OK, prev.text
            html = prev.json()["html"]
            assert html and "John Doe" in html, html
            assert "{{" not in html, f"unfilled token in preview: {html}"

            # 5) And it drives the PDF (via fpdf2's HTML engine — no LibreOffice).
            pdf = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/pdf")
            assert pdf.status_code == status.HTTP_200_OK
            assert pdf.content[:4] == b"%PDF"

            # 6) Re-open the wizard: saved choices are pre-filled as suggestions.
            reopened = c.get("/api/v1/enterprise/settings/payslip/document/mapping")
            assert reopened.status_code == status.HTTP_200_OK, reopened.text
            again = {s["label"]: s["suggested_token"] for s in reopened.json()["slots"]}
            assert again.get("Basic Salary") == "amount.BASIC"
        finally:
            c.delete("/api/v1/enterprise/settings/payslip/document")


def test_payslip_pdf_uses_template_content(monkeypatch: pytest.MonkeyPatch) -> None:
    """The /pdf endpoint must render the company's template (not the built-in
    layout) when it's enabled. We disable PDF compression so the template's text
    is greppable in the raw bytes, proving which renderer ran."""
    import io as _io

    from docx import Document as _Doc

    from app.services import pdf_service

    _orig = pdf_service._PayslipPDF.__init__

    def _init(self: Any, *a: Any, **k: Any) -> None:
        _orig(self, *a, **k)
        self.set_compression(False)

    monkeypatch.setattr(pdf_service._PayslipPDF, "__init__", _init)

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def blank() -> bytes:
        d = _Doc()
        d.add_paragraph("ZZUNIQUEHEADER ACME")
        t = d.add_table(rows=2, cols=4)
        t.cell(0, 0).text = "Basic Salary"; t.cell(0, 2).text = "Provident Fund"
        t.cell(1, 0).text = "Net Pay"; t.cell(1, 2).text = "Total Deductions"
        b = _io.BytesIO(); d.save(b); return b.getvalue()

    with authed_client() as c:
        slip = _paid_payslip(c)
        try:
            scan = c.post(
                "/api/v1/enterprise/settings/payslip/document/scan",
                files={"file": ("blank.docx", blank(), mime)},
            )
            slots = scan.json()["slots"]
            mapping = {
                str(s["index"]): s["suggested_token"]
                for s in slots
                if s["suggested_token"]
            }
            c.put(
                "/api/v1/enterprise/settings/payslip/document/mapping",
                json={"mapping": mapping},
            )
            pdf = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/pdf").content
            assert b"ZZUNIQUEHEADER" in pdf, "PDF did not use the template content"
            assert b"Net Payable" not in pdf, "PDF used the built-in layout"
            # Downloads must not be cached, so re-downloading reflects new state.
            r = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/pdf")
            assert "no-store" in r.headers.get("cache-control", "")
        finally:
            c.delete("/api/v1/enterprise/settings/payslip/document")


def test_payslip_mapping_detects_text_value_and_vertical_cells() -> None:
    """Employee/company detail cells — which hold sample TEXT, or sit in a
    label-above-value stack — must be detected and filled, not just blank/numeric
    amount cells."""
    import io as _io

    from docx import Document as _Doc

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def template() -> bytes:
        d = _Doc()
        # Employee details with SAMPLE TEXT values (horizontal pairs).
        det = d.add_table(rows=1, cols=4)
        det.cell(0, 0).text = "Employee Name"; det.cell(0, 1).text = "SAMPLE PERSON"
        det.cell(0, 2).text = "Designation"; det.cell(0, 3).text = "Sample Role"
        # Vertical block: labels above blank values.
        ver = d.add_table(rows=2, cols=2)
        ver.cell(0, 0).text = "PAN"; ver.cell(0, 1).text = "Net Pay"
        ver.cell(1, 0).text = ""; ver.cell(1, 1).text = ""
        b = _io.BytesIO(); d.save(b); return b.getvalue()

    with authed_client() as c:
        slip = _paid_payslip(c)
        try:
            scan = c.post(
                "/api/v1/enterprise/settings/payslip/document/scan",
                files={"file": ("emp.docx", template(), mime)},
            )
            labels = {s["label"]: s["suggested_token"] for s in scan.json()["slots"]}
            assert labels.get("Employee Name") == "employee.name"
            assert labels.get("Designation") == "employee.designation"
            assert labels.get("PAN") == "employee.pan"  # vertical block detected
            assert labels.get("Net Pay") == "net"

            mapping = {
                str(s["index"]): s["suggested_token"]
                for s in scan.json()["slots"]
                if s["suggested_token"]
            }
            c.put(
                "/api/v1/enterprise/settings/payslip/document/mapping",
                json={"mapping": mapping},
            )
            prev = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/preview-html")
            html = prev.json()["html"]
            assert "John Doe" in html, html          # employee name filled
            assert "SAMPLE PERSON" not in html         # sample text overwritten
        finally:
            c.delete("/api/v1/enterprise/settings/payslip/document")


def test_payslip_mapping_detects_label_fields_inside_table_cells() -> None:
    """Header fields like 'Company Name: ____' that live INSIDE table cells
    (single-column or merged rows) must be detected and filled — doc.paragraphs
    doesn't include cell paragraphs, so this guards that path."""
    import io as _io

    from docx import Document as _Doc

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def tmpl() -> bytes:
        d = _Doc()
        hdr = d.add_table(rows=2, cols=1)  # single-column header block
        hdr.cell(0, 0).text = "Company Name: " + "_" * 20
        hdr.cell(1, 0).text = "Employee Name: " + "_" * 20
        merged = d.add_table(rows=1, cols=2)
        mc = merged.cell(0, 0).merge(merged.cell(0, 1))  # merged-across row
        mc.text = "Designation: " + "_" * 20
        b = _io.BytesIO(); d.save(b); return b.getvalue()

    with authed_client() as c:
        slip = _paid_payslip(c)
        try:
            scan = c.post(
                "/api/v1/enterprise/settings/payslip/document/scan",
                files={"file": ("t.docx", tmpl(), mime)},
            )
            slots = scan.json()["slots"]
            by_label = {s["label"]: s["suggested_token"] for s in slots}
            assert by_label.get("Company Name") == "company_name"
            assert by_label.get("Employee Name") == "employee.name"
            assert by_label.get("Designation") == "employee.designation"

            mapping = {
                str(s["index"]): s["suggested_token"] for s in slots if s["suggested_token"]
            }
            c.put(
                "/api/v1/enterprise/settings/payslip/document/mapping",
                json={"mapping": mapping},
            )
            html = c.get(
                f"/api/v1/enterprise/payroll/payslips/{slip['id']}/preview-html"
            ).json()["html"]
            assert "John Doe" in html, html  # employee name filled inside the cell
        finally:
            c.delete("/api/v1/enterprise/settings/payslip/document")


def test_payslip_mapping_handles_bracketed_logo_placeholder() -> None:
    """A standalone "[ COMPANY LOGO ]" placeholder is detected, auto-suggested to
    the logo field, and replaced in the generated payslip (not left as-is)."""
    import io as _io

    from docx import Document as _Doc

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def tmpl() -> bytes:
        d = _Doc()
        d.add_paragraph("[ COMPANY LOGO ]")
        d.add_paragraph("Employee Name: ")
        b = _io.BytesIO(); d.save(b); return b.getvalue()

    with authed_client() as c:
        slip = _paid_payslip(c)
        try:
            scan = c.post(
                "/api/v1/enterprise/settings/payslip/document/scan",
                files={"file": ("t.docx", tmpl(), mime)},
            )
            slots = scan.json()["slots"]
            logo = next(s for s in slots if s["label"] == "COMPANY LOGO")
            assert logo["suggested_token"] == "logo"
            mapping = {
                str(s["index"]): s["suggested_token"] for s in slots if s["suggested_token"]
            }
            c.put(
                "/api/v1/enterprise/settings/payslip/document/mapping",
                json={"mapping": mapping},
            )
            html = c.get(
                f"/api/v1/enterprise/payroll/payslips/{slip['id']}/preview-html"
            ).json()["html"]
            assert "[ COMPANY LOGO ]" not in html, html  # placeholder replaced
        finally:
            c.delete("/api/v1/enterprise/settings/payslip/document")


def test_amount_to_words_conversion() -> None:
    """Net-pay-in-words spelling: Indian lakh/crore grouping, paise, currency
    units, rounding carry-over, and unknown-code fallback."""
    from app.services.docx_service import amount_to_words as w

    assert w(12345.50, "INR") == (
        "Rupees Twelve Thousand Three Hundred Forty-Five and Fifty Paise Only"
    )
    assert w(1234567, "INR") == (
        "Rupees Twelve Lakh Thirty-Four Thousand Five Hundred Sixty-Seven Only"
    )
    assert w(10000000, "INR") == "Rupees One Crore Only"
    assert w(999.999, "INR") == "Rupees One Thousand Only"  # rounds up
    assert w(0, "INR") == "Rupees Zero Only"
    assert w(1234567.05, "USD") == (
        "Dollars One Million Two Hundred Thirty-Four Thousand "
        "Five Hundred Sixty-Seven and Five Cents Only"
    )
    assert w(1000, "XYZ") == "XYZ One Thousand Only"  # unknown code falls back


def test_payslip_net_pay_in_words_mapped_and_filled() -> None:
    """A "Net Pay in words" line is detected, auto-suggested to net_in_words, and
    filled with the spelled-out net pay in the generated payslip."""
    import io as _io

    from docx import Document as _Doc

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def tmpl() -> bytes:
        d = _Doc()
        d.add_paragraph("Net Pay: {{ net }}")
        d.add_paragraph("Net Pay in words: ")
        b = _io.BytesIO(); d.save(b); return b.getvalue()

    with authed_client() as c:
        slip = _paid_payslip(c)
        try:
            scan = c.post(
                "/api/v1/enterprise/settings/payslip/document/scan",
                files={"file": ("t.docx", tmpl(), mime)},
            )
            slots = scan.json()["slots"]
            words_slot = next(s for s in slots if s["label"] == "Net Pay in words")
            assert words_slot["suggested_token"] == "net_in_words"
            mapping = {
                str(s["index"]): s["suggested_token"] for s in slots if s["suggested_token"]
            }
            c.put(
                "/api/v1/enterprise/settings/payslip/document/mapping",
                json={"mapping": mapping},
            )
            html = c.get(
                f"/api/v1/enterprise/payroll/payslips/{slip['id']}/preview-html"
            ).json()["html"]
            assert "Net Pay in words:" in html
            # The spelled-out amount (whatever the net is) ends with "Only".
            assert "Rupees" in html and "Only" in html, html
        finally:
            c.delete("/api/v1/enterprise/settings/payslip/document")


def test_payslip_logo_lands_on_placeholder_not_stray_token() -> None:
    """The logo image embeds AT the "[ COMPANY LOGO ]" placeholder (keeping its
    alignment), and any stray {{ logo }} token elsewhere is cleared — so the
    logo is never duplicated or dropped in a random spot."""
    import io as _io

    from docx import Document as _Doc
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from PIL import Image

    from app.services import docx_service

    png = _io.BytesIO(); Image.new("RGB", (120, 50), (200, 0, 0)).save(png, "PNG")

    d = _Doc()
    d.add_paragraph("ACME PAYSLIP")
    p = d.add_paragraph("[ COMPANY LOGO ]"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph("Stray: {{ logo }}")  # mis-placed token must not win
    b = _io.BytesIO(); d.save(b); tpl = b.getvalue()

    out = docx_service.render_payslip_docx(tpl, {"logo": "ACME"}, logo_image=png.getvalue())
    rd = _Doc(_io.BytesIO(out))

    def img_count(par) -> int:
        return len(par._p.findall(".//" + qn("a:blip")))

    placeholder = rd.paragraphs[1]
    assert img_count(placeholder) == 1, "logo not at the [ COMPANY LOGO ] placeholder"
    assert placeholder.alignment == WD_ALIGN_PARAGRAPH.CENTER  # alignment preserved
    assert sum(img_count(par) for par in rd.paragraphs) == 1, "logo duplicated"
    assert not any("{{" in par.text for par in rd.paragraphs), "stray token left"


def test_payslip_logo_image_used_from_settings(monkeypatch: pytest.MonkeyPatch) -> None:
    """When a branding logo_url is set, the [ COMPANY LOGO ] placeholder renders
    the actual logo IMAGE (no re-asking) — embedded as <img> in the preview."""
    import io as _io

    from docx import Document as _Doc
    from PIL import Image

    from app.services import docx_service

    # Avoid any network: the logo "fetch" returns a small in-memory PNG.
    png = _io.BytesIO(); Image.new("RGB", (120, 50), (10, 80, 200)).save(png, "PNG")
    monkeypatch.setattr(docx_service, "fetch_logo_image", lambda url: png.getvalue())

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def tmpl() -> bytes:
        d = _Doc()
        d.add_paragraph("[ COMPANY LOGO ]")
        b = _io.BytesIO(); d.save(b); return b.getvalue()

    with authed_client() as c:
        slip = _paid_payslip(c)
        try:
            # Configure a branding logo URL in settings.
            c.put(
                "/api/v1/enterprise/settings/payslip",
                json={"logo_url": "https://example.com/logo.png"},
            )
            scan = c.post(
                "/api/v1/enterprise/settings/payslip/document/scan",
                files={"file": ("t.docx", tmpl(), mime)},
            )
            slots = scan.json()["slots"]
            mapping = {
                str(s["index"]): s["suggested_token"] for s in slots if s["suggested_token"]
            }
            c.put(
                "/api/v1/enterprise/settings/payslip/document/mapping",
                json={"mapping": mapping},
            )
            html = c.get(
                f"/api/v1/enterprise/payroll/payslips/{slip['id']}/preview-html"
            ).json()["html"]
            assert '<img src="data:image' in html, "logo image not embedded"
        finally:
            c.delete("/api/v1/enterprise/settings/payslip/document")
            c.put("/api/v1/enterprise/settings/payslip", json={"logo_url": None})


def test_payslip_mapping_reads_existing_placeholders() -> None:
    """If the uploaded template already contains {{ tokens }}, the wizard pre-maps
    those slots from the existing token (not just the label)."""
    import io as _io

    from docx import Document as _Doc

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def tokenised() -> bytes:
        d = _Doc()
        d.add_paragraph("Staff Member: {{ employee.name }}")  # label we don't know
        t = d.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "Take Home"; t.cell(0, 1).text = "{{ net }}"
        b = _io.BytesIO(); d.save(b); return b.getvalue()

    with authed_client() as c:
        try:
            scan = c.post(
                "/api/v1/enterprise/settings/payslip/document/scan",
                files={"file": ("t.docx", tokenised(), mime)},
            )
            by_label = {s["label"]: s["suggested_token"] for s in scan.json()["slots"]}
            # Pre-mapped from the existing placeholders, despite unknown labels.
            assert by_label.get("Staff Member") == "employee.name"
            assert by_label.get("Take Home") == "net"
        finally:
            c.delete("/api/v1/enterprise/settings/payslip/document")


def test_payslip_doc_has_tokens_flag() -> None:
    """A raw (token-free) upload reports doc_has_tokens=False; after mapping it
    flips to True. This is what the UI uses to warn 'no data will fill'."""
    import io as _io

    from docx import Document as _Doc

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def blank() -> bytes:
        d = _Doc()
        d.add_paragraph("Net Pay: ")
        t = d.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "Basic Salary"
        b = _io.BytesIO(); d.save(b); return b.getvalue()

    with authed_client() as c:
        try:
            # Direct upload of a token-free doc -> no fillable fields.
            up = c.put(
                "/api/v1/enterprise/settings/payslip/document",
                files={"file": ("blank.docx", blank(), mime)},
            )
            assert up.status_code == status.HTTP_200_OK, up.text
            assert up.json()["has_doc_template"] is True
            assert up.json()["doc_has_tokens"] is False

            # Map it via the wizard -> tokens injected, flag flips.
            scan = c.post(
                "/api/v1/enterprise/settings/payslip/document/scan",
                files={"file": ("blank.docx", blank(), mime)},
            )
            slots = scan.json()["slots"]
            mapping = {
                str(s["index"]): s["suggested_token"] for s in slots if s["suggested_token"]
            }
            applied = c.put(
                "/api/v1/enterprise/settings/payslip/document/mapping",
                json={"mapping": mapping},
            )
            assert applied.json()["doc_has_tokens"] is True
        finally:
            c.delete("/api/v1/enterprise/settings/payslip/document")


def test_payslip_preview_html_null_without_template() -> None:
    with authed_client() as c:
        slip = _paid_payslip(c)
        c.delete("/api/v1/enterprise/settings/payslip/document")
        r = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/preview-html")
        assert r.status_code == status.HTTP_200_OK
        assert r.json()["html"] is None


def test_payslip_mapping_requires_scanned_original() -> None:
    with authed_client() as c:
        c.delete("/api/v1/enterprise/settings/payslip/document")
        # No scanned original -> 404.
        r = c.put(
            "/api/v1/enterprise/settings/payslip/document/mapping",
            json={"mapping": {"0": "net"}},
        )
        assert r.status_code == status.HTTP_404_NOT_FOUND


def _configure_smtp(monkeypatch: pytest.MonkeyPatch) -> None:
    from app.services import email_service

    monkeypatch.setattr(email_service._settings, "smtp_host", "smtp.test")
    monkeypatch.setattr(email_service._settings, "smtp_port", 587)
    monkeypatch.setattr(email_service._settings, "smtp_use_ssl", False)
    monkeypatch.setattr(email_service._settings, "smtp_username", "payroll@test.dev")
    monkeypatch.setattr(email_service._settings, "smtp_password", "app-password")
    monkeypatch.setattr(email_service._settings, "smtp_from_email", "payroll@test.dev")


class _FakeSMTP:
    """Stand-in for smtplib.SMTP capturing the sent message (no network)."""

    sent: list[Any] = []

    def __init__(self, host: str, port: int) -> None:
        self.host = host
        self.port = port

    def __enter__(self) -> "_FakeSMTP":
        return self

    def __exit__(self, *exc: object) -> None:
        return None

    def starttls(self, context: object = None) -> None:
        pass

    def login(self, user: str, password: str) -> None:
        self.user = user

    def send_message(self, msg: Any) -> None:
        _FakeSMTP.sent.append(msg)


def test_email_payslip_not_configured(monkeypatch: pytest.MonkeyPatch) -> None:
    from app.services import email_service

    monkeypatch.setattr(email_service._settings, "smtp_username", "")
    monkeypatch.setattr(email_service._settings, "smtp_password", "")
    monkeypatch.setattr(email_service._settings, "smtp_from_email", "")
    with authed_client() as c:
        slip = _paid_payslip(c)
        r = c.post(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/email")
        assert r.status_code == status.HTTP_503_SERVICE_UNAVAILABLE


def test_email_payslip_success(monkeypatch: pytest.MonkeyPatch) -> None:
    from app.services import email_service

    _configure_smtp(monkeypatch)
    _FakeSMTP.sent = []
    monkeypatch.setattr(email_service.smtplib, "SMTP", _FakeSMTP)

    with authed_client() as c:
        slip = _paid_payslip(c)
        r = c.post(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/email")
        assert r.status_code == status.HTTP_200_OK, r.text
        body = r.json()
        assert body["sent"] is True
        assert body["to"] == "john@example.com"
        # One message was sent, addressed correctly, with the PDF attached.
        assert len(_FakeSMTP.sent) == 1
        msg = _FakeSMTP.sent[0]
        assert msg["To"] == "john@example.com"
        attachments = [p for p in msg.iter_attachments()]
        assert len(attachments) == 1
        assert attachments[0].get_filename().endswith(".pdf")
        assert attachments[0].get_content_type() == "application/pdf"


def test_email_payslip_requires_paid_cycle() -> None:
    """A payslip from a non-PAID cycle is gated (403) for email too."""
    with authed_client() as c:
        _create_structure(c, employee_id=EMP1)
        cid = _create_cycle(c)["id"]
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run")  # -> PROCESSING
    # Fetch the hidden payslip id directly from the DB.
    conn = psycopg2.connect(_DSN)
    conn.autocommit = True
    cur = conn.cursor()
    cur.execute("SELECT id FROM payslips WHERE cycle_id = %s;", (cid,))
    slip_id = cur.fetchone()[0]
    cur.close()
    conn.close()
    with authed_client() as c:
        assert (
            c.post(f"/api/v1/enterprise/payroll/payslips/{slip_id}/email").status_code
            == status.HTTP_403_FORBIDDEN
        )


def test_viewer_cannot_email_payslip() -> None:
    with authed_client(ADMIN) as admin:
        slip = _paid_payslip(admin)
    with authed_client(VIEWER) as viewer:
        r = viewer.post(f"/api/v1/enterprise/payroll/payslips/{slip['id']}/email")
        assert r.status_code == status.HTTP_403_FORBIDDEN


# ---------------------------------------------------------------------------
# Statutory compliance — Phase 1 (PF / ESI / PT)
# ---------------------------------------------------------------------------
def test_statutory_off_by_default_unchanged() -> None:
    """With statutory toggles off, the calculation is identical to before."""
    res = compute_payslip(_make_struct(), Decimal("0"), Decimal("30"))
    assert res["gross_earnings"] == Decimal("71000.00")
    assert res["total_deductions"] == Decimal("1800.00")  # only the manual PF line
    assert res["employer_contributions"] == []


def test_statutory_pf_capped_at_ceiling() -> None:
    """PF (employee) = 12% of min(basic, 15000); employer splits into EPS+EPF."""
    struct = _make_struct(
        components=[{"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 40000}],
        default_deductions=[],
        pf_enabled=True,
        pf_cap_at_ceiling=True,
    )
    res = compute_payslip(struct, Decimal("0"), Decimal("30"))
    pf = next(d for d in res["deductions"] if d["code"] == "PF")
    assert pf["amount"] == 1800.0  # 12% of 15000
    eps = next(c for c in res["employer_contributions"] if c["code"] == "EPS_ER")
    epf = next(c for c in res["employer_contributions"] if c["code"] == "PF_ER")
    assert eps["amount"] == 1250.0  # 8.33% of 15000
    assert epf["amount"] == 550.0   # 12% - 8.33%
    assert res["statutory"]["pf"]["wage_considered"] == 15000.0


def test_statutory_pf_uncapped() -> None:
    struct = _make_struct(
        components=[{"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 40000}],
        default_deductions=[],
        pf_enabled=True,
        pf_cap_at_ceiling=False,
    )
    res = compute_payslip(struct, Decimal("0"), Decimal("30"))
    pf = next(d for d in res["deductions"] if d["code"] == "PF")
    assert pf["amount"] == 4800.0  # 12% of 40000


def test_statutory_esi_excluded_above_limit() -> None:
    """ESI only applies when gross <= 21000; a high earner is not covered."""
    struct = _make_struct(
        components=[{"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 40000}],
        default_deductions=[],
        esi_enabled=True,
    )
    res = compute_payslip(struct, Decimal("0"), Decimal("30"))
    assert all(d["code"] != "ESI" for d in res["deductions"])
    assert res["statutory"]["esi"]["covered"] is False


def test_statutory_esi_applies_within_limit() -> None:
    struct = _make_struct(
        components=[{"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 20000}],
        default_deductions=[],
        esi_enabled=True,
    )
    res = compute_payslip(struct, Decimal("0"), Decimal("30"))
    esi = next(d for d in res["deductions"] if d["code"] == "ESI")
    assert esi["amount"] == 150.0  # 0.75% of 20000
    esi_er = next(c for c in res["employer_contributions"] if c["code"] == "ESI_ER")
    assert esi_er["amount"] == 650.0  # 3.25% of 20000


def test_statutory_pt_by_state() -> None:
    struct = _make_struct(
        components=[{"code": "BASIC", "label": "Basic", "type": "fixed", "amount": 40000}],
        default_deductions=[],
        pt_enabled=True,
    )
    res_ka = compute_payslip(struct, Decimal("0"), Decimal("30"), pt_state="KA")
    pt = next(d for d in res_ka["deductions"] if d["code"] == "PT")
    assert pt["amount"] == 200.0
    # Unknown/unset state -> no PT, with a note recorded.
    res_none = compute_payslip(struct, Decimal("0"), Decimal("30"), pt_state=None)
    assert all(d["code"] != "PT" for d in res_none["deductions"])
    assert res_none["statutory"]["pt"]["amount"] == 0.0


def test_statutory_run_persists_employer_and_totals() -> None:
    """Enabling statutory on a structure flows through run -> payslip + totals."""
    with authed_client() as c:
        # Put EMP1 in Karnataka so PT applies.
        c.put(f"/api/v1/enterprise/employees/{EMP1}", json={"state": "KA"})
        _create_structure(
            c,
            employee_id=EMP1,
            default_deductions=[],  # statutory replaces manual deduction lines
            pf_enabled=True,
            esi_enabled=True,
            pt_enabled=True,
        )
        cid = _create_cycle(c)["id"]
        assert c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/run").status_code == status.HTTP_200_OK

        got = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}").json()
        # gross 56000 (BASIC 40000 + HRA 16000); ESI excluded (>21000);
        # PF employee 1800 + PT 200 -> net 54000; employer PF 1800.
        assert got["totals"]["gross"] == 56000.0
        assert got["totals"]["net"] == 54000.0
        assert got["totals"]["employer_cost"] == 1800.0
        assert got["totals"]["total_cost"] == 57800.0

        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/approve")
        c.post(f"/api/v1/enterprise/payroll/cycles/{cid}/mark-paid")
        slip = c.get(f"/api/v1/enterprise/payroll/cycles/{cid}/payslips").json()[0]
        detail = c.get(f"/api/v1/enterprise/payroll/payslips/{slip['id']}").json()
        ded_codes = {d["code"] for d in detail["deductions"]}
        assert {"PF", "PT"}.issubset(ded_codes)
        er_codes = {c["code"] for c in detail["employer_contributions"]}
        assert {"PF_ER", "EPS_ER"}.issubset(er_codes)
