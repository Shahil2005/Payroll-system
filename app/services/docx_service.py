"""Payslip generation from an uploaded .docx template.

Two parts:
- ``render_payslip_docx`` fills an uploaded Word template (docxtpl / Jinja2
  tokens) with a payslip context. Pure-Python, always available.
- ``docx_to_pdf`` converts the filled .docx to PDF using LibreOffice headless
  (or, on Windows with MS Word, the ``docx2pdf`` package). This needs a
  converter installed on the host; when none is found it returns ``None`` and
  the caller falls back to the built-in fpdf2 layout.

Token reference for template authors (Jinja2 syntax, e.g. ``{{ net_pay }}``):
    company_name, ref, status, currency, period_start, period_end, pay_date,
    employee.{name,email,code,pan,uan}, company.{name,legal_name,pan,tan},
    gross, total_deductions, net (formatted strings),
    earnings / deductions / employer_contributions: lists of
        {code, label, amount}  -- use a docxtpl table-row loop
    lop_days, paid_days, working_days
"""
from __future__ import annotations

import io
import os
import shutil
import subprocess
import tempfile
from typing import Any

from docxtpl import DocxTemplate

# OOXML (.docx) files are ZIP archives — they start with the local-file magic.
_ZIP_MAGIC = b"PK\x03\x04"


def looks_like_docx(data: bytes) -> bool:
    return data[:4] == _ZIP_MAGIC


def render_payslip_docx(template_bytes: bytes, context: dict[str, Any]) -> bytes:
    """Fill the uploaded .docx template with ``context`` and return .docx bytes."""
    doc = DocxTemplate(io.BytesIO(template_bytes))
    doc.render(context)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def sample_payslip_template() -> bytes:
    """A ready-to-use .docx payslip template with correct tokens.

    Built with python-docx so every token sits in a single run (no Word
    run-splitting), which is the usual reason hand-typed tokens don't fill.
    Admins download this, restyle it, and re-upload. Uses the preformatted
    ``*_lines`` tokens (reliable) rather than docxtpl's finicky ``{%tr%}`` loops.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    d = Document()
    title = d.add_paragraph()
    run = title.add_run("{{ company_name }}")
    run.bold = True
    run.font.size = Pt(18)
    sub = d.add_paragraph("PAYSLIP — Ref #{{ ref }}  ·  Status: {{ status }}")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    d.add_paragraph("Employee: {{ employee.name }}  ({{ employee.code }})")
    d.add_paragraph("Email: {{ employee.email }}")
    d.add_paragraph("Pay period: {{ period_start }} to {{ period_end }}   ·   Pay date: {{ pay_date }}")

    d.add_paragraph("")
    h1 = d.add_paragraph().add_run("Earnings")
    h1.bold = True
    d.add_paragraph("{{ earnings_lines }}")
    d.add_paragraph("Gross Earnings\t{{ gross }}")

    d.add_paragraph("")
    h2 = d.add_paragraph().add_run("Deductions")
    h2.bold = True
    d.add_paragraph("{{ deductions_lines }}")
    d.add_paragraph("Total Deductions\t{{ total_deductions }}")

    d.add_paragraph("")
    h3 = d.add_paragraph().add_run("Attendance")
    h3.bold = True
    d.add_paragraph("Working days: {{ working_days }}   LOP: {{ lop_days }}   Paid: {{ paid_days }}")

    d.add_paragraph("")
    net = d.add_paragraph()
    nrun = net.add_run("NET PAYABLE\t{{ net }}")
    nrun.bold = True
    nrun.font.size = Pt(14)

    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _find_soffice() -> str | None:
    """Locate LibreOffice: explicit override, then PATH, then common installs."""
    override = os.getenv("PAYROLL_SOFFICE_PATH")
    if override and os.path.exists(override):
        return override
    for name in ("soffice", "soffice.exe", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    for guess in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/opt/libreoffice/program/soffice",
    ):
        if os.path.exists(guess):
            return guess
    return None


def pdf_conversion_available() -> bool:
    """True if a docx->pdf converter (LibreOffice or MS Word) is on this host."""
    if _find_soffice():
        return True
    try:  # MS Word via docx2pdf (Windows/macOS, optional dependency)
        import docx2pdf  # noqa: F401

        return True
    except Exception:
        return False


def docx_to_pdf(docx_bytes: bytes) -> bytes | None:
    """Convert .docx bytes to PDF. Returns None when conversion is unavailable
    or fails, so the caller can fall back to the built-in PDF layout."""
    soffice = _find_soffice()
    if soffice:
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "payslip.docx")
            with open(src, "wb") as fh:
                fh.write(docx_bytes)
            # A per-call user profile avoids lock clashes under concurrency.
            profile = os.path.join(tmp, "profile")
            try:
                subprocess.run(
                    [
                        soffice,
                        f"-env:UserInstallation=file:///{profile.replace(os.sep, '/')}",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        tmp,
                        src,
                    ],
                    check=True,
                    capture_output=True,
                    timeout=90,
                )
            except (subprocess.SubprocessError, OSError):
                return None
            pdf_path = os.path.join(tmp, "payslip.pdf")
            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as fh:
                    return fh.read()
        return None

    # Fallback: MS Word via docx2pdf, if installed.
    try:
        import docx2pdf
    except Exception:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "payslip.docx")
        dst = os.path.join(tmp, "payslip.pdf")
        with open(src, "wb") as fh:
            fh.write(docx_bytes)
        try:
            docx2pdf.convert(src, dst)
        except Exception:
            return None
        if os.path.exists(dst):
            with open(dst, "rb") as fh:
                return fh.read()
    return None
